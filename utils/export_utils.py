import os
from io import BytesIO
from docx import Document
from fpdf import FPDF

# For Gujarati/Hindi support in PDF, we need a custom font.
# We will stub the font path. In a real scenario, a TTF file like "NotoSansGujarati-Regular.ttf" must be provided.
FONT_PATH = os.path.join(os.path.dirname(__file__), "fonts", "NotoSansGujarati.ttf")

def create_docx(text, title="Generated Summary"):
    """
    Creates a Word (.docx) document in memory and returns it as a BytesIO object.
    Word handles unicode (Gujarati/Hindi) natively.
    """
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def create_pdf(text, title="Generated Summary", language="en"):
    """
    Creates a PDF document in memory and returns it as a BytesIO object.
    Note: fpdf2 requires a custom TTF font for non-Latin characters.
    """
    pdf = FPDF()
    pdf.add_page()
    
    # Try to load a Unicode font if it's not English to support Gujarati/Hindi. 
    # If font is missing, fpdf will fallback or error.
    try:
        if language != "en" and os.path.exists(FONT_PATH):
            pdf.add_font("NotoSans", style="", fname=FONT_PATH, uni=True)
            pdf.set_font("NotoSans", size=12)
        else:
            # Default Arial for English
            pdf.set_font("Arial", size=12)
    except Exception as e:
        print(f"Font loading error: {e}. Falling back to default.")
        pdf.set_font("Arial", size=12)
            
    # Add Title
    pdf.set_font(pdf.font_family, style="B", size=16)
    pdf.cell(200, 10, txt=title, ln=True, align="C")
    pdf.ln(10)
    
    # Add Body
    pdf.set_font(pdf.font_family, style="", size=12)
    
    # fpdf2 multi_cell handles text wrapping
    # encode to utf-8 if using default font, though FPDF2 handles strings mostly ok.
    pdf.multi_cell(0, 10, txt=text)
    
    bio = BytesIO(pdf.output(dest='S'))
    bio.seek(0)
    return bio
